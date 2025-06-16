
# # you need to install all these in your terminal
# # pip install streamlit
# # pip install scikit-learn
# # pip install python-docx
# # pip install PyPDF2


# import streamlit as st
# import pickle
# import docx  # Extract text from Word file
# import PyPDF2  # Extract text from PDF
# import re
# import gdown
# file_url = "https://drive.google.com/file/d/1vQi73uuWL1X-UPumNjbf1363GElghKxY/view?usp=sharing"
# gdown.download(file_url, "clf.pkl", quiet=False)
# # Load pre-trained model and TF-IDF vectorizer (ensure these are saved earlier)
# svc_model = pickle.load(open('clf.pkl', 'rb'))  # Example file name, adjust as needed
# tfidf = pickle.load(open('tfidf.pkl', 'rb'))  # Example file name, adjust as needed
# le = pickle.load(open('encoder.pkl', 'rb'))  # Example file name, adjust as needed


# # Function to clean resume text
# def cleanResume(txt):
#     cleanText = re.sub('http\S+\s', ' ', txt)
#     cleanText = re.sub('RT|cc', ' ', cleanText)
#     cleanText = re.sub('#\S+\s', ' ', cleanText)
#     cleanText = re.sub('@\S+', '  ', cleanText)
#     cleanText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', cleanText)
#     cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
#     cleanText = re.sub('\s+', ' ', cleanText)
#     return cleanText


# # Function to extract text from PDF
# def extract_text_from_pdf(file):
#     pdf_reader = PyPDF2.PdfReader(file)
#     text = ''
#     for page in pdf_reader.pages:
#         text += page.extract_text()
#     return text


# # Function to extract text from DOCX
# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     text = ''
#     for paragraph in doc.paragraphs:
#         text += paragraph.text + '\n'
#     return text


# # Function to extract text from TXT with explicit encoding handling
# def extract_text_from_txt(file):
#     # Try using utf-8 encoding for reading the text file
#     try:
#         text = file.read().decode('utf-8')
#     except UnicodeDecodeError:
#         # In case utf-8 fails, try 'latin-1' encoding as a fallback
#         text = file.read().decode('latin-1')
#     return text


# # Function to handle file upload and extraction
# def handle_file_upload(uploaded_file):
#     file_extension = uploaded_file.name.split('.')[-1].lower()
#     if file_extension == 'pdf':
#         text = extract_text_from_pdf(uploaded_file)
#     elif file_extension == 'docx':
#         text = extract_text_from_docx(uploaded_file)
#     elif file_extension == 'txt':
#         text = extract_text_from_txt(uploaded_file)
#     else:
#         raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
#     return text


# # Function to predict the category of a resume
# def pred(input_resume):
#     # Preprocess the input text (e.g., cleaning, etc.)
#     cleaned_text = cleanResume(input_resume)

#     # Vectorize the cleaned text using the same TF-IDF vectorizer used during training
#     vectorized_text = tfidf.transform([cleaned_text])

#     # Convert sparse matrix to dense
#     vectorized_text = vectorized_text.toarray()

#     # Prediction
#     predicted_category = svc_model.predict(vectorized_text)

#     # get name of predicted category
#     predicted_category_name = le.inverse_transform(predicted_category)

#     return predicted_category_name[0]  # Return the category name


# # Streamlit app layout
# def main():
#     st.set_page_config(page_title="Resume Category Prediction", page_icon="📄", layout="wide")

#     st.title("Resume Category Prediction App")
#     st.markdown("Upload a resume in PDF, TXT, or DOCX format and get the predicted job category.")

#     # File upload section
#     uploaded_file = st.file_uploader("Upload a Resume", type=["pdf", "docx", "txt"])

#     if uploaded_file is not None:
#         # Extract text from the uploaded file
#         try:
#             resume_text = handle_file_upload(uploaded_file)
#             st.write("Successfully extracted the text from the uploaded resume.")

#             # Display extracted text (optional)
#             if st.checkbox("Show extracted text", False):
#                 st.text_area("Extracted Resume Text", resume_text, height=300)

#             # Make prediction
#             st.subheader("Predicted Category")
#             category = pred(resume_text)
#             st.write(f"The predicted category of the uploaded resume is: **{category}**")

#         except Exception as e:
#             st.error(f"Error processing the file: {str(e)}")


# if __name__ == "__main__":
# # you need to install all these in your terminal
# # pip install streamlit
# # pip install scikit-learn
# # pip install python-docx
# # pip install PyPDF2


#     import streamlit as st
#     import pickle
#     import docx  # Extract text from Word file
#     import PyPDF2  # Extract text from PDF
#     import re
    
#     # Load pre-trained model and TF-IDF vectorizer (ensure these are saved earlier)
#     svc_model = pickle.load(open('clf.pkl', 'rb'))  # Example file name, adjust as needed
#     tfidf = pickle.load(open('tfidf.pkl', 'rb'))  # Example file name, adjust as needed
#     le = pickle.load(open('encoder.pkl', 'rb'))  # Example file name, adjust as needed
    
    
#     # Function to clean resume text
#     def cleanResume(txt):
#         cleanText = re.sub('http\S+\s', ' ', txt)
#         cleanText = re.sub('RT|cc', ' ', cleanText)
#         cleanText = re.sub('#\S+\s', ' ', cleanText)
#         cleanText = re.sub('@\S+', '  ', cleanText)
#         cleanText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', cleanText)
#         cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
#         cleanText = re.sub('\s+', ' ', cleanText)
#         return cleanText
    
    
#     # Function to extract text from PDF
#     def extract_text_from_pdf(file):
#         pdf_reader = PyPDF2.PdfReader(file)
#         text = ''
#         for page in pdf_reader.pages:
#             text += page.extract_text()
#         return text
    
    
#     # Function to extract text from DOCX
#     def extract_text_from_docx(file):
#         doc = docx.Document(file)
#         text = ''
#         for paragraph in doc.paragraphs:
#             text += paragraph.text + '\n'
#         return text
    
    
#     # Function to extract text from TXT with explicit encoding handling
#     def extract_text_from_txt(file):
#         # Try using utf-8 encoding for reading the text file
#         try:
#             text = file.read().decode('utf-8')
#         except UnicodeDecodeError:
#             # In case utf-8 fails, try 'latin-1' encoding as a fallback
#             text = file.read().decode('latin-1')
#         return text
    
    
#     # Function to handle file upload and extraction
#     def handle_file_upload(uploaded_file):
#         file_extension = uploaded_file.name.split('.')[-1].lower()
#         if file_extension == 'pdf':
#             text = extract_text_from_pdf(uploaded_file)
#         elif file_extension == 'docx':
#             text = extract_text_from_docx(uploaded_file)
#         elif file_extension == 'txt':
#             text = extract_text_from_txt(uploaded_file)
#         else:
#             raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")
#         return text
    
    
#     # Function to predict the category of a resume
#     def pred(input_resume):
#         # Preprocess the input text (e.g., cleaning, etc.)
#         cleaned_text = cleanResume(input_resume)
    
#         # Vectorize the cleaned text using the same TF-IDF vectorizer used during training
#         vectorized_text = tfidf.transform([cleaned_text])
    
#         # Convert sparse matrix to dense
#         vectorized_text = vectorized_text.toarray()
    
#         # Prediction
#         predicted_category = svc_model.predict(vectorized_text)
    
#         # get name of predicted category
#         predicted_category_name = le.inverse_transform(predicted_category)
    
#         return predicted_category_name[0]  # Return the category name
    
    
#     # Streamlit app layout
#     def main():
#         st.set_page_config(page_title="Resume Category Prediction", page_icon="📄", layout="wide")
    
#         st.title("Resume Category Prediction App")
#         st.markdown("Upload a resume in PDF, TXT, or DOCX format and get the predicted job category.")
    
#         # File upload section
#         uploaded_file = st.file_uploader("Upload a Resume", type=["pdf", "docx", "txt"])
    
#         if uploaded_file is not None:
#             # Extract text from the uploaded file
#             try:
#                 resume_text = handle_file_upload(uploaded_file)
#                 st.write("Successfully extracted the text from the uploaded resume.")
    
#                 # Display extracted text (optional)
#                 if st.checkbox("Show extracted text", False):
#                     st.text_area("Extracted Resume Text", resume_text, height=300)
    
#                 # Make prediction
#                 st.subheader("Predicted Category")
#                 category = pred(resume_text)
#                 st.write(f"The predicted category of the uploaded resume is: **{category}**")
    
#             except Exception as e:
#                 st.error(f"Error processing the file: {str(e)}")


# if __name__ == "__main__":
#     main()







# requirements.txt should include:
# streamlit==1.35.0
# scikit-learn==1.5.2
# python-docx==1.1.0
# PyPDF2==3.0.1
# gdown==5.1.0
# numpy==1.26.4
# pandas==2.1.3

import streamlit as st
import pickle
import docx
import PyPDF2
import re
import os
import gdown
from io import BytesIO

# --- Model Loading ---
@st.cache_resource
def load_models():
    try:
        # 1. Download clf.pkl from Google Drive if missing
        if not os.path.exists('clf.pkl'):
            gdown.download(
                'https://drive.google.com/uc?id=1vQi73uuWL1X-UPumNjbf1363GElghKxY',
                'clf.pkl',
                quiet=False
            )
        
        # 2. Load all models
        return {
            'svc_model': pickle.load(open('clf.pkl', 'rb')),
            'tfidf': pickle.load(open('tfidf.pkl', 'rb')),
            'encoder': pickle.load(open('encoder.pkl', 'rb'))
        }
    except Exception as e:
        st.error(f"❌ Model loading failed: {str(e)}")
        st.error("Required files:")
        st.error("1. clf.pkl (from Google Drive)")
        st.error("2. tfidf.pkl (in GitHub repo)")
        st.error("3. encoder.pkl (in GitHub repo)")
        st.stop()

# --- Text Cleaning ---
def cleanResume(txt):
    cleanText = re.sub('http\S+\s', ' ', txt)
    cleanText = re.sub('RT|cc', ' ', cleanText)
    cleanText = re.sub('#\S+\s', ' ', cleanText)
    cleanText = re.sub('@\S+', '  ', cleanText)
    cleanText = re.sub('[%s]' % re.escape("""!"#$%&'()*+,-./:;<=>?@[\]^_`{|}~"""), ' ', cleanText)
    cleanText = re.sub(r'[^\x00-\x7f]', ' ', cleanText)
    cleanText = re.sub('\s+', ' ', cleanText)
    return cleanText

# --- File Extractors ---
def extract_text_from_pdf(file):
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
        return ''.join([page.extract_text() or '' for page in pdf_reader.pages])
    except Exception as e:
        st.error(f"PDF error: {str(e)}")
        return ""

def extract_text_from_docx(file):
    try:
        doc = docx.Document(BytesIO(file.read()))
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"DOCX error: {str(e)}")
        return ""

def extract_text_from_txt(file):
    try:
        file.seek(0)
        return str(file.read(), 'utf-8', errors='ignore')
    except Exception as e:
        st.error(f"TXT error: {str(e)}")
        return ""

# --- Main App ---
def main():
    st.set_page_config(page_title="Resume Classifier", layout="wide")
    st.title("📄 Resume Category Prediction")
    
    # Load models
    models = load_models()
    
    # File upload
    uploaded_file = st.file_uploader("Upload Resume (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"])
    
    if uploaded_file:
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Resume Content")
            try:
                # Extract text
                if uploaded_file.name.endswith('.pdf'):
                    text = extract_text_from_pdf(uploaded_file)
                elif uploaded_file.name.endswith('.docx'):
                    text = extract_text_from_docx(uploaded_file)
                else:
                    text = extract_text_from_txt(uploaded_file)
                
                # Show extracted text
                if st.toggle("Show raw text"):
                    st.text_area("Content", text, height=300)
                
                # Process if text exists
                if text.strip():
                    with col2:
                        st.subheader("Prediction")
                        cleaned = cleanResume(text)
                        vector = models['tfidf'].transform([cleaned]).toarray()
                        pred = models['encoder'].inverse_transform(
                            models['svc_model'].predict(vector)
                        )[0]
                        st.success(f"**Predicted Category:** {pred}")
                else:
                    st.warning("No text could be extracted")
                    
            except Exception as e:
                st.error(f"❌ Processing error: {str(e)}")

if __name__ == "__main__":
    main()
